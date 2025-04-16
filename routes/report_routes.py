from flask import Blueprint, render_template, request, send_file, flash, make_response
import os
import csv
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
from database.models import Equipment
from .auth_routes import login_required, role_required  # Use relative import

report_bp = Blueprint('report_bp', __name__)

@report_bp.route('/generate_report', methods=['GET'])
def generate_report():
    """Generate a CSV report of all equipment."""
    report_folder = os.path.join('static', 'uploads', 'reports')
    os.makedirs(report_folder, exist_ok=True)
    report_path = os.path.join(report_folder, 'equipment_report.csv')

    # Fetch all equipment data
    equipment_list = Equipment.query.all()

    # Write data to CSV
    with open(report_path, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(['ID', 'Name', 'Type', 'Status', 'Location', 'Photo', 'Manual', 'QR Code'])
        for equipment in equipment_list:
            writer.writerow([
                equipment.id,
                equipment.name,
                equipment.type,
                equipment.status,
                equipment.location,
                equipment.photo,
                equipment.manual,
                equipment.qr_code
            ])

    return send_file(report_path, as_attachment=True, download_name='equipment_report.csv')

@report_bp.route('/reports_page', methods=['GET'])
def reports_page():
    # Retrieve filter parameters
    selected_type = request.args.get('type', '')
    location_filter = request.args.get('location', '')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    # Build the query dynamically based on filters
    query = Equipment.query
    if selected_type:
        query = query.filter(Equipment.type == selected_type)
    if location_filter:
        query = query.filter(Equipment.location.contains(location_filter))
    if date_from:
        query = query.filter(Equipment.activation_date >= date_from)
    if date_to:
        query = query.filter(Equipment.activation_date <= date_to)

    # Fetch filtered equipment
    equipment_list = query.all()
    return render_template('reports.html', equipment_list=equipment_list, selected_type=selected_type, location_filter=location_filter, date_from=date_from, date_to=date_to)

def get_filtered_equipment(args):
    """Retrieve filtered equipment based on request arguments."""
    query = Equipment.query

    # Apply filters based on request arguments
    if args.get('type'):
        query = query.filter(Equipment.type == args.get('type'))
    if args.get('location'):
        query = query.filter(Equipment.location.contains(args.get('location')))
    if args.get('date_from'):
        query = query.filter(Equipment.activation_date >= args.get('date_from'))
    if args.get('date_to'):
        query = query.filter(Equipment.activation_date <= args.get('date_to'))

    return query.all()

@report_bp.route('/export_word', methods=['GET'])
def export_word():
    # Retrieve filtered equipment
    equipment_list = get_filtered_equipment(request.args)

    # Create Word document
    doc = Document()
    doc.add_heading('Сүзілген жабдықтар тізімі', level=1)
    for equipment in equipment_list:
        doc.add_paragraph(f"Атауы: {equipment.name}")
        doc.add_paragraph(f"Түрі: {equipment.type}")
        doc.add_paragraph(f"Мәртебесі: {equipment.status}")
        doc.add_paragraph(f"Орналасуы: {equipment.location}")
        doc.add_paragraph("-" * 20)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Return Word file as response
    response = make_response(buffer.read())
    response.headers['Content-Disposition'] = 'attachment; filename=filtered_report.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return response

@report_bp.route('/export_full_word', methods=['GET'])
def export_full_word():
    # Retrieve all equipment
    equipment_list = Equipment.query.all()

    # Create Word document
    doc = Document()
    current_date = datetime.now().strftime('%Y-%m-%d')  # Get the current date
    doc.add_heading(f'Толық жабдықтар тізімі ({current_date})', level=1)  # Add the date to the heading

    for equipment in equipment_list:
        doc.add_paragraph(f"Атауы: {equipment.name}")
        doc.add_paragraph(f"Түрі: {equipment.type}")
        doc.add_paragraph(f"Мәртебесі: {equipment.status}")
        doc.add_paragraph(f"Орналасуы: {equipment.location}")
        doc.add_paragraph(f"Түгендеу нөмірі: {equipment.inventory_number}")
        doc.add_paragraph(f"Іске қосу күні: {equipment.activation_date}")
        doc.add_paragraph("-" * 20)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Return Word file as response
    response = make_response(buffer.read())
    response.headers['Content-Disposition'] = f'attachment; filename=full_report_{current_date}.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return response
